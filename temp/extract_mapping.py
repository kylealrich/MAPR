"""Extract Field Mapping section from the Cerner GLTrans Analysis and Design Specification."""
import sys
from docx import Document
from docx.table import Table

doc_path = r'input/Aultman Health_I8_Cerner_GLTrans_Analysis and Design Specification_v1.1.docx'
doc = Document(doc_path)

# First, print all headings to understand document structure
print("=" * 80)
print("DOCUMENT HEADINGS")
print("=" * 80)
for i, para in enumerate(doc.paragraphs):
    if para.style.name.startswith('Heading'):
        print(f"  [{i}] {para.style.name}: {para.text}")

# Print all tables with context
print("\n" + "=" * 80)
print("ALL TABLES IN DOCUMENT")
print("=" * 80)
for t_idx, table in enumerate(doc.tables):
    # Get the paragraph just before this table for context
    print(f"\n--- Table {t_idx + 1} ---")
    # Print header row
    if table.rows:
        header_cells = [cell.text.strip() for cell in table.rows[0].cells]
        print(f"  Headers: {header_cells}")
        print(f"  Row count: {len(table.rows)}")
        # Print first 3 data rows as sample
        for r_idx, row in enumerate(table.rows[1:4], 1):
            cells = [cell.text.strip() for cell in row.cells]
            print(f"  Row {r_idx}: {cells}")
        if len(table.rows) > 4:
            print(f"  ... ({len(table.rows) - 4} more rows)")
